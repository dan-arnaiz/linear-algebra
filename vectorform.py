from docx import Document

# Create a new Document
doc = Document()

# Add the formulas
doc.add_paragraph("1. Norm, Dot Product and Distance in R")
doc.add_paragraph("- Norm (Length of a vector): ||v|| = sqrt(v1² + v2² + ... + vn²)")
doc.add_paragraph("- Dot Product: u.v = u1v1 + u2v2 + ... + unvn")
doc.add_paragraph("- Distance: d(u, v) = ||u - v||")

doc.add_paragraph("2. Orthogonality")
doc.add_paragraph("- Two vectors are orthogonal if their dot product is zero: u.v = 0")

doc.add_paragraph("3. Cross Product (only in R³)")
doc.add_paragraph("- u x v = (u2v3 - u3v2, u3v1 - u1v3, u1v2 - u2v1)")

doc.add_paragraph("4. Eigenvalues and Eigenvectors")
doc.add_paragraph("- Av = λv, where A is a matrix, v is an eigenvector, and λ is an eigenvalue")
doc.add_paragraph("- Eigenvalues are solutions to the characteristic equation: det(A - λI) = 0")

doc.add_paragraph("5. Diagonalization")
doc.add_paragraph("- A matrix A is diagonalizable if there exists an invertible matrix P such that P⁻¹AP is a diagonal matrix D. The columns of P are the eigenvectors of A.")

doc.add_paragraph("6. Complex Vector Spaces")
doc.add_paragraph("- These are vector spaces where the scalars are complex numbers. All the rules of real vector spaces apply, but with complex arithmetic.")

doc.add_paragraph("7. Inner Products")
doc.add_paragraph("- An inner product on a vector space V is a function that associates a number, denoted as ⟨u, v⟩, with each pair of vectors u and v of V. It satisfies certain properties like commutativity, linearity in the first argument, and positive-definiteness.")

doc.add_paragraph("8. Angle and Orthogonality in Inner Product Spaces")
doc.add_paragraph("- Angle: cos θ = ⟨u, v⟩ / (||u|| ||v||)")
doc.add_paragraph("- Orthogonality: Two vectors are orthogonal if their inner product is zero: ⟨u, v⟩ = 0")

# Save the document
doc.save("vector_formulas.docx")